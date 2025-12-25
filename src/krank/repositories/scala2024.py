import pandas as pd
from docx import Document
from pooch import Unzip

from ._base import KrankRepo


class Scala2024(KrankRepo):
    """
    Tressoldi, Patrizio (2024). Dreaming similar contents intentionally. figshare. Dataset. https://doi.org/10.6084/m9.figshare.25488214.v1
    Dreams transcriptions related to the study "Dreaming similar contents intentionally: an empirical evidence", available as a preprint at https://osf.io/preprints/psyarxiv/9r3z5
    Scala, V. F., Cozzi, G., Castiglioni, M. L., Finelli, M. E., Vitali, P., Liberale, L., … Tressoldi, P. (2024, March 28). Dreaming similar contents intentionally: an empirical evidence. https://doi.org/10.31234/osf.io/9r3z5
    """
    def __init__(self):
        super().__init__(repo_id="scala2024")

    def read_dreams(self):
        fp = self.pup.fetch("Dreams.zip", processor=Unzip())
        records = []
        for path in fp:
            pstem = Path(path).stem
            matches = re.match(r"^AI(\w+)Series(Control)?\w+(\d)$", pstem)
            series_number = matches.group(1)
            is_control = matches.group(2)
            session_number = matches.group(3)
            document = Document(path)
            all_text = "\n".join(p.text for p in document.paragraphs)
            for found in re.findall(r"\n([A-Z][A-Za-z]+)(\d_?\d?): ([\w\W]*?)(?=\n\n)", all_text):
                author, number, dream = found
                records.append(dict(
                    series_number=series_number, is_control=is_control, session_number=session_number,
                    author=author, number=number, dream=dream,
                    )
                )
            # for para in document.paragraphs:
            #     if (text := para.text):
            #         re.findall(r"\n([A-Z][A-Za-z]+)(\d_?\d?): ([\w\W]*?)(?=\n\n)", all_text)
            #         matches = re.match(r"^(\w+)([\d_]+): (.*)", text, re.MULTILINE)
            #         author = 
        dreams = pd.DataFrame(records)
        dtypes = {
            "series_number": "string",
            "is_control": "bool",
            "session_number": "int",
            "author": "string",
            "number": "string",
            "dream": "string",
        }
        dreams = dreams.astype(dtypes)
        dreams["series_number"] = pd.Categorical(dreams["series_number"], ordered=True)
        dreams["author"] = pd.Categorical(dreams["author"], ordered=False)
        dreams["session_number"] = pd.Categorical(dreams["session_number"], ordered=True)
        return dreams

